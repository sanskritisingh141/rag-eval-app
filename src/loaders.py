from langchain_community.document_loaders import PyPDFLoader, TextLoader
from docx import Document
from langchain_core.documents import Document as LCDocument
import tempfile


def load_pdf(uploaded_file):
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
        tmp.write(uploaded_file.read())
        tmp_path = tmp.name

    loader = PyPDFLoader(tmp_path)
    return loader.load()


def load_txt(uploaded_file):
    text = uploaded_file.read().decode("utf-8")
    return [LCDocument(page_content=text, metadata={"source": uploaded_file.name})]


def load_docx(uploaded_file):
    doc = Document(uploaded_file)
    text = "\n".join([p.text for p in doc.paragraphs])
    return [LCDocument(page_content=text, metadata={"source": uploaded_file.name})]


def load_document(uploaded_file):
    name = uploaded_file.name.lower()

    if name.endswith(".pdf"):
        return load_pdf(uploaded_file)

    if name.endswith(".txt"):
        return load_txt(uploaded_file)

    if name.endswith(".docx"):
        return load_docx(uploaded_file)

    raise ValueError("Unsupported file type")